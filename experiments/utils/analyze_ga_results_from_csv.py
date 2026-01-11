#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
analyze_ga_results_from_csv.py

Analysis of GA strategies (Diversity vs Adaptive / Random) using a CSV input:
- Adjusted Fitness Improvement (AFI) using avg_fitness_after_first
- Corrected REC improvement (lower is better for minimization)
- AUCC improvement (lower is better)
- Generations improvement (lower is better)
- Stability as absolute ΔSD (diversity_SD - baseline_SD)

Inputs:
    A CSV with columns:
        experiment, number_of_variables, saturation_generations, strategy,
        avg_min_fitness, avg_fitness_after_first, sd_min_fitness,
        avg_generations, relative_early_convergence, area_under_convergence_curve

Outputs (written to --outdir):
- rel_improvements_per_combination_AFI.csv
- rel_improvements_overall_AFI.csv
- rel_improvements_standard_by_dim_sat_AFI.csv
- rel_improvements_custom_by_sat_AFI.csv
- GA_comparison_AFI_REC_AUCC.docx (requires python-docx)

Usage:
    python analyze_ga_results_from_csv.py --input _final_results_merged.csv --outdir outputs
"""
import argparse, os, sys
import math

import numpy as np
import pandas as pd

def adjusted_fitness_improvement(f_first, f_base, f_div):
    """AFI (%) = ((f_base - f_div) / (f_first - f_div)) * 100"""
    denom = f_first - f_div
    if denom == 0:
        return np.nan
    return ((f_base - f_div) / denom) * 100

def _sd_spc(base_sd, div_sd, eps=1e-12):
    den = base_sd + div_sd
    if np.isnan(base_sd) or np.isnan(div_sd) or den < eps:
        return np.nan
    return 200.0 * (base_sd - div_sd) / den
# -----------------------------
# Plotting Utilities (matplotlib only)
# -----------------------------
import matplotlib.pyplot as plt

def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path

def plot_lines(df, x_col, y_col, hue_col, title, xlabel, ylabel, outpath):
    """Simple line chart with one figure, multiple lines by hue_col (e.g., compare_to)."""
    plt.figure()
    for key in sorted(df[hue_col].dropna().unique()):
        sub = df[df[hue_col] == key].sort_values(x_col)
        plt.plot(sub[x_col].astype(int), sub[y_col], marker='o', label=str(key))
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.legend()
    plt.tight_layout()
    plt.savefig(outpath, dpi=200)
    plt.close()


def compute_pairwise(df):
    """Compute AFI, REC (corrected), AUCC, generations, ΔSD for Diversity vs baselines per (exp, dim, sat)."""
    rows = []
    group_cols = ["experiment","number_of_variables","saturation_generations"]
    required = {
        "strategy","avg_min_fitness","avg_fitness_after_first","sd_min_fitness",
        "avg_generations","relative_early_convergence","area_under_convergence_curve"
    }
    miss = required - set(df.columns)
    if miss:
        raise KeyError(f"Input file is missing required columns: {sorted(miss)}")

    for (exp, nvars, sat), sub in df.groupby(group_cols):
        means = sub.groupby("strategy").agg(
            fmin=("avg_min_fitness","mean"),
            ffirst=("avg_fitness_after_first","mean"),
            sd=("sd_min_fitness","mean"),
            gens=("avg_generations","mean"),
            rec=("relative_early_convergence","mean"),
            aucc=("area_under_convergence_curve","mean")
        )
        if "diversity mutation" not in means.index:
            continue  # need diversity to compare against
        div = means.loc["diversity mutation"]

        for baseline in ["adaptive mutation","random mutation"]:
            if baseline not in means.index:
                continue
            base = means.loc[baseline]

            afi = adjusted_fitness_improvement(div.ffirst, base.fmin, div.fmin)
            aucc_imp = np.nan if (base.aucc==0 or np.isnan(base.aucc) or np.isnan(div.aucc)) else 100*(base.aucc - div.aucc)/abs(base.aucc)
            rec_imp  = np.nan if (base.rec==0  or np.isnan(base.rec)  or np.isnan(div.rec))  else 100*(base.rec - div.rec)/abs(base.rec)  # lower is better
            gens_imp = np.nan if (base.gens==0 or np.isnan(base.gens) or np.isnan(div.gens)) else 100*(base.gens - div.gens)/abs(base.gens)
            dsd = _sd_spc(base.sd, div.sd)
            if math.isnan(dsd):
                pass
            rows.append({
                "experiment": exp,
                "number_of_variables": nvars,
                "saturation": sat,
                "compare_to": baseline,
                "adjusted_fitness_improvement_%": afi,
                "aucc_improvement_%": aucc_imp,
                "rec_improvement_%": rec_imp,
                "generations_improvement_%": gens_imp,
                "stability_delta_SD": dsd
            })
    return pd.DataFrame(rows)

def save_standard_plots(std_agg, outdir):
    charts_dir = _ensure_dir(os.path.join(outdir, "charts_standard"))
    # Metrics to visualize
    metrics = [
        ("adjusted_fitness_improvement_%", "AFI (%)"),
        ("rec_improvement_%", "REC Δ% (higher = faster early convergence)"),
        ("aucc_improvement_%", "AUCC Δ%"),
        ("generations_improvement_%", "Generations Δ%"),
        ("stability_delta_SD", "ΔSD (abs; negative = Diversity more stable)")
    ]
    for nvars in sorted(std_agg["number_of_variables"].dropna().unique()):
        sub_dim = std_agg[std_agg["number_of_variables"] == nvars]
        for metric, label in metrics:
            title = f"Standard — {int(nvars)} vars: {label} vs Saturation"
            fname = f"standard_{int(nvars)}vars_{metric.replace('%','pct').replace(' ','_')}.png"
            outpath = os.path.join(charts_dir, fname)
            plot_lines(sub_dim, x_col="saturation", y_col=metric, hue_col="compare_to",
                       title=title, xlabel="Saturation (no-improvement generations)", ylabel=label, outpath=outpath)

def save_custom_plots(cust_agg, outdir):
    charts_dir = _ensure_dir(os.path.join(outdir, "charts_custom"))
    metrics = [
        ("adjusted_fitness_improvement_%", "AFI (%)"),
        ("rec_improvement_%", "REC Δ% (higher = faster early convergence)"),
        ("aucc_improvement_%", "AUCC Δ%"),
        ("generations_improvement_%", "Generations Δ%"),
        ("stability_delta_SD", "ΔSD (abs; negative = Diversity more stable)")
    ]
    for metric, label in metrics:
        title = f"Custom (pooled): {label} vs Saturation"
        fname = f"custom_{metric.replace('%','pct').replace(' ', '_')}.png"
        outpath = os.path.join(charts_dir, fname)
        plot_lines(cust_agg, x_col="saturation", y_col=metric, hue_col="compare_to",
                   title=title, xlabel="Saturation (no-improvement generations)", ylabel=label, outpath=outpath)

def write_word_report(overall_df, std_agg, cust_agg, outpath):
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Skipping Word report.", file=sys.stderr)
        return

    doc = Document()
    doc.add_heading("GA Strategy Comparison with AFI (Adjusted Fitness), Corrected REC & AUCC", level=1)
    doc.add_paragraph(
        "AFI accounts for the improvement margin from first-generation average fitness down to the diversity minimum. "
        "REC is treated as a minimization metric (lower REC = faster early convergence); positive REC Δ% indicates "
        "Diversity converges earlier. AUCC & Generations are lower-is-better. Stability is SD Δ% - positive SD Δ% indicates lower standard deviation of the diversity mutation, which means the higher fitness stability. "
        "(negative → Diversity more stable)."
    )

    # Overall
    doc.add_heading("Overall (all functions, all settings)", level=2)
    t = doc.add_table(rows=1, cols=6)
    hdr = t.rows[0].cells
    hdr[0].text = "Compared To"; hdr[1].text = "AFI (%)"; hdr[2].text = "AUCC Δ%"; hdr[3].text = "REC Δ%"; hdr[4].text = "Generations Δ%"; hdr[5].text = "SD Δ%"
    for _, r in overall_df.iterrows():
        rw = t.add_row().cells
        rw[0].text = str(r["compare_to"])
        rw[1].text = f"{r['adjusted_fitness_improvement_%']:.2f}"
        rw[2].text = f"{r['aucc_improvement_%']:.2f}"
        rw[3].text = f"{r['rec_improvement_%']:.2f}"
        rw[4].text = f"{r['generations_improvement_%']:.2f}"
        rw[5].text = f"{r['stability_delta_SD']:.4f}"

    # Standard Benchmarks (by dimension & saturation)
    doc.add_heading("Standard Benchmarks (by dimension & saturation)", level=2)
    for nvars in [2,3,5,7]:
        sub = std_agg[std_agg["number_of_variables"]==nvars]
        doc.add_heading(f"{nvars} variables", level=3)
        if sub.empty:
            doc.add_paragraph("No data."); continue
        t = doc.add_table(rows=1, cols=7)
        hdr = t.rows[0].cells
        hdr[0].text = "Compared To"; hdr[1].text = "Saturation"; hdr[2].text = "AFI (%)"; hdr[3].text = "AUCC Δ%"; hdr[4].text = "REC Δ%"; hdr[5].text = "Generations Δ%"; hdr[6].text = "SD Δ%"
        for _, r in sub.iterrows():
            rw = t.add_row().cells
            rw[0].text = r["compare_to"]
            rw[1].text = str(int(r["saturation"]))
            rw[2].text = f"{r['adjusted_fitness_improvement_%']:.2f}"
            rw[3].text = f"{r['aucc_improvement_%']:.2f}"
            rw[4].text = f"{r['rec_improvement_%']:.2f}"
            rw[5].text = f"{r['generations_improvement_%']:.2f}"
            rw[6].text = f"{r['stability_delta_SD']:.4f}"

    # Custom Trigonometric Functions (by saturation only, pooled)
    doc.add_heading("Custom Trigonometric Functions (pooled by saturation)", level=2)
    t = doc.add_table(rows=1, cols=7)
    hdr = t.rows[0].cells
    hdr[0].text = "Compared To"; hdr[1].text = "Saturation"; hdr[2].text = "AFI (%)"; hdr[3].text = "AUCC Δ%"; hdr[4].text = "REC Δ%"; hdr[5].text = "Generations Δ%"; hdr[6].text = "SD Δ%"
    for _, r in cust_agg.iterrows():
        rw = t.add_row().cells
        rw[0].text = r["compare_to"]
        rw[1].text = str(int(r["saturation"]))
        rw[2].text = f"{r['adjusted_fitness_improvement_%']:.2f}"
        rw[3].text = f"{r['aucc_improvement_%']:.2f}"
        rw[4].text = f"{r['rec_improvement_%']:.2f}"
        rw[5].text = f"{r['generations_improvement_%']:.2f}"
        rw[6].text = f"{r['stability_delta_SD']:.4f}"

    doc.save(outpath)

def generate_results(input_path: str, outdir: str):
    os.makedirs(outdir, exist_ok=True)
    df = pd.read_csv(input_path)

    # Pairwise comparisons
    rel_all = compute_pairwise(df)

    # Save per-combination
    per_combo_path = os.path.join(outdir, "rel_improvements_per_combination_AFI.csv")
    rel_all.to_csv(per_combo_path, index=False)

    # Overall averages
    overall = (rel_all.groupby("compare_to")[
                   ["adjusted_fitness_improvement_%", "aucc_improvement_%", "rec_improvement_%",
                    "generations_improvement_%", "stability_delta_SD"]
               ].mean().reset_index())
    overall_path = os.path.join(outdir, "rel_improvements_overall_AFI.csv")
    overall.to_csv(overall_path, index=False)

    # Standard vs Custom aggregations
    standard_funcs = [
        "Sphere Function", "Rosenbrock Function", "Ackley Function", "Rastrigin Function",
        "Griewank Function", "Beale Function", "Himmelblau Function", "Booth Function",
        "Styblinski Tang Function"
    ]
    custom_funcs = [
        "Highly Coupled Trigonometric Function",
        "Moderately Coupled Trigonometric Function",
        "Separated Trig Arithmetic Function"
    ]

    rel_std = rel_all[rel_all["experiment"].isin(standard_funcs) & rel_all["number_of_variables"].isin([2, 3, 5, 7])]
    std_agg = (rel_std.groupby(["number_of_variables", "saturation", "compare_to"])[
                   ["adjusted_fitness_improvement_%", "aucc_improvement_%", "rec_improvement_%",
                    "generations_improvement_%", "stability_delta_SD"]
               ].mean().reset_index().sort_values(["number_of_variables", "compare_to", "saturation"]))
    std_agg_path = os.path.join(outdir, "rel_improvements_standard_by_dim_sat_AFI.csv")
    std_agg.to_csv(std_agg_path, index=False)

    rel_cust = rel_all[rel_all["experiment"].isin(custom_funcs)]
    cust_agg = (rel_cust.groupby(["saturation", "compare_to"])[
                    ["adjusted_fitness_improvement_%", "aucc_improvement_%", "rec_improvement_%",
                     "generations_improvement_%", "stability_delta_SD"]
                ].mean().reset_index().sort_values(["compare_to", "saturation"]))
    cust_agg_path = os.path.join(outdir, "rel_improvements_custom_by_sat_AFI.csv")
    cust_agg.to_csv(cust_agg_path, index=False)

    # Word report
    # Charts (no seaborn; one chart per figure)
    try:
        save_standard_plots(std_agg, outdir)
        save_custom_plots(cust_agg, outdir)
    except Exception as e:
        print("Warning: could not generate plots:", e, file=sys.stderr)

    word_path = os.path.join(outdir, "GA_comparison_AFI_REC_AUCC.docx")
    try:
        write_word_report(overall, std_agg, cust_agg, word_path)
    except Exception as e:
        print("Warning: could not write Word report:", e, file=sys.stderr)

    print("Saved:")
    print(" -", per_combo_path)
    print(" -", overall_path)
    print(" -", std_agg_path)
    print(" -", cust_agg_path)
    print(" -", word_path)


def main():
    parser = argparse.ArgumentParser(description="Analyze GA results from CSV with AFI/REC/AUCC.")
    parser.add_argument("--input", default="_final_results_merged.csv", help="Path to GA results CSV.")
    parser.add_argument("--outdir", default=".", help="Output directory for CSVs and DOCX.")
    args = parser.parse_args()
    generate_results(args.input, args.outdir)

if __name__ == "__main__":
    main()






