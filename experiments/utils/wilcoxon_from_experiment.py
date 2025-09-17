#!/usr/bin/env python3
import argparse, re
from pathlib import Path
import numpy as np
import pandas as pd
from scipy.stats import wilcoxon, mannwhitneyu

try:
    from docx import Document
except Exception:
    Document = None

# ---------- Helpers ----------
ALIASES = {
    "experiment": ["experiment"],
    "strategy": ["strategy","method","mutation","algo","algorithm"],
    "run_id": ["run_id","run","seed","trial"],
    "fitness": ["fitness","cost","objective","best_fitness","best_cost","value"],
}

def pick_col(df, keys):
    cols = {c.lower(): c for c in df.columns}
    for k in keys:
        if k in cols:
            return cols[k]
    return None

def std_strategy(s: str) -> str:
    s = str(s).strip().lower()
    if "diversity" in s:
        return "GAdapt-Diversity"
    if "adaptive" in s:
        return "PyGAD-Adaptive"
    if "random" in s:
        return "PyGAD-Random"
    return s or "Unknown"

# Format: "beale_func (2 variables, saturation = 10)"
EXP_RE = re.compile(
    r"^\s*(?P<func>[^()]+?)\s*\(\s*(?P<vars>\d+)\s*variables?\s*,\s*saturation\s*=\s*(?P<sat>\d+)\s*\)\s*$",
    re.I
)
def parse_experiment(exp_str: str):
    m = EXP_RE.match(str(exp_str))
    if not m:
        return None, None, None
    return m.group("func").strip(), int(m.group("vars")), int(m.group("sat"))

def cliff_delta(x, y):
    x = np.asarray(x); y = np.asarray(y)
    nx, ny = len(x), len(y)
    if nx == 0 or ny == 0: return np.nan
    xx = np.repeat(x[:, None], ny, axis=1)
    yy = np.repeat(y[None, :], nx, axis=0)
    return (np.sum(xx < yy) - np.sum(xx > yy)) / (nx * ny)

def rank_biserial_paired(x, y):
    diff = np.asarray(y) - np.asarray(x)
    wins = int(np.sum(diff > 0))
    losses = int(np.sum(diff < 0))
    return (wins - losses) / (wins + losses) if (wins + losses) > 0 else 0.0

# ---------- Core pipeline ----------
def collect_csvs(root: Path):
    return [p for p in root.rglob("*.csv")]

def build_per_run_minima(csv_paths):
    rows, skipped = [], []
    for p in csv_paths:
        try:
            df = pd.read_csv(p)
        except Exception as e:
            skipped.append((str(p), f"read_error: {e}"))
            continue

        exp_col = pick_col(df, ALIASES["experiment"])
        strat_col = pick_col(df, ALIASES["strategy"])
        run_col = pick_col(df, ALIASES["run_id"]) or "run_id"
        fit_col = pick_col(df, ALIASES["fitness"])

        if exp_col is None or strat_col is None or fit_col is None:
            skipped.append((str(p), "missing required columns"))
            continue

        if run_col not in df.columns:
            df[run_col] = 0

        func, n_vars, saturation = parse_experiment(df[exp_col].iloc[0])
        if func is None:
            skipped.append((str(p), f"could_not_parse_experiment: {df[exp_col].iloc[0]!r}"))
            continue

        strategy_std = std_strategy(df[strat_col].iloc[0])

        grp = df.groupby(run_col)[fit_col].min().reset_index().rename(columns={fit_col: "best_fitness_final"})
        for _, r in grp.iterrows():
            rows.append({
                "function": func,
                "n_vars": n_vars,
                "saturation": saturation,
                "strategy": strategy_std,
                "run_id": int(r[run_col]) if pd.notna(r[run_col]) else 0,
                "best_fitness_final": float(r["best_fitness_final"]),
                "__source": str(p)
            })
    return pd.DataFrame(rows), skipped

def run_statistics(per_run_min: pd.DataFrame):
    stats_rows = []
    for (func, n, sat), g in per_run_min.groupby(["function","n_vars","saturation"]):
        for other in ["PyGAD-Random","PyGAD-Adaptive"]:
            A = g[g["strategy"]=="GAdapt-Diversity"]
            B = g[g["strategy"]==other]
            if A.empty or B.empty:
                continue
            common = sorted(set(A["run_id"]).intersection(set(B["run_id"])))
            if len(common) >= 5:
                x = A.set_index("run_id").loc[common]["best_fitness_final"].values
                y = B.set_index("run_id").loc[common]["best_fitness_final"].values
                stat, p = wilcoxon(x, y, alternative="less")
                effect = rank_biserial_paired(x, y)
                test, N = "Wilcoxon (paired)", len(common)
                medA, medB = np.median(x), np.median(y)
            elif len(A)>=3 and len(B)>=3:
                x = A["best_fitness_final"].values; y = B["best_fitness_final"].values
                stat, p = mannwhitneyu(x, y, alternative="less")
                effect = cliff_delta(x, y)
                test, N = "Mann–Whitney (unpaired)", min(len(x), len(y))
                medA, medB = np.median(x), np.median(y)
            else:
                stat=p=effect=medA=medB=np.nan; test="Insufficient"; N=0
            stats_rows.append({
                "function": func, "n_vars": n, "saturation": sat,
                "comparison": f"GAdapt-Diversity vs {other}",
                "test": test, "N": N, "stat": stat, "p_value": p, "effect_size": effect,
                "median_GAdapt": medA, "median_PyGAD": medB, "median_diff": medB - medA
            })
    return pd.DataFrame(stats_rows)

def export_word_table(stats_df: pd.DataFrame, docx_path: Path):
    if Document is None:
        print("python-docx not installed; skipping Word export.")
        return
    headers = ["Function","Dim","Sat","Comparison","Test","N","Stat","p-value","Effect","Med. diff"]
    doc = Document()
    doc.add_heading("Wilcoxon / Mann–Whitney Results", level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    for i,h in enumerate(headers):
        table.rows[0].cells[i].text = h

    def fmt(x):
        if pd.isna(x): return ""
        try: return f"{float(x):.3g}"
        except: return str(x)

    for _, r in stats_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r.get("function",""))
        cells[1].text = str(r.get("n_vars",""))
        cells[2].text = str(r.get("saturation",""))
        cells[3].text = str(r.get("comparison",""))
        cells[4].text = str(r.get("test",""))
        cells[5].text = str(r.get("N",""))
        cells[6].text = fmt(r.get("stat"))
        cells[7].text = fmt(r.get("p_value"))
        cells[8].text = fmt(r.get("effect_size"))
        cells[9].text = fmt(r.get("median_diff"))
    doc.save(docx_path)

# ---------- Main ----------
def main():
    ap = argparse.ArgumentParser(description="Wilcoxon/Mann–Whitney analysis from folder of per-run CSVs with experiment column.")
    ap.add_argument("--folder", required=True, type=Path, help="Folder containing unpacked CSVs.")
    ap.add_argument("--out", required=True, type=Path, help="Output directory.")
    args = ap.parse_args()

    out_dir = args.out; out_dir.mkdir(parents=True, exist_ok=True)
    csvs = collect_csvs(args.folder)

    per_run_min, skipped = build_per_run_minima(csvs)
    per_run_min.to_csv(out_dir/"per_run_minima_from_experiment.csv", index=False)

    stats_df = run_statistics(per_run_min)
    stats_df.to_csv(out_dir/"wilcoxon_results_from_experiment.csv", index=False)
    export_word_table(stats_df, out_dir/"wilcoxon_results_from_experiment.docx")

    if skipped:
        pd.DataFrame(skipped, columns=["file","reason"]).to_csv(out_dir/"skipped_files.csv", index=False)
        print(f"Some files skipped; see skipped_files.csv")

if __name__ == "__main__":
    main()
