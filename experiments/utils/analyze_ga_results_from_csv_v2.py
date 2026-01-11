#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
analyze_ga_results_from_csv_v2.py

Analysis of GA strategies (Diversity vs Adaptive / Random) using a CSV input:
- Adjusted Fitness Improvement (AFI) using avg_fitness_after_first
- Corrected REC improvement (lower is better for minimization; higher Δ% means Diversity better)
- AUCC improvement (lower is better; higher Δ% means Diversity better)
- Generations improvement (lower is better; higher Δ% means Diversity faster)
  *Reported in 2 ways: (a) classic slice-weighted; (b) function-equal (each function counts once)*
- Stability as SD Δ% = 100*(SD_base - SD_div)/SD_base (positive ⇒ Diversity more stable)

Inputs:
    CSV columns:
        experiment, number_of_variables, saturation_generations, strategy,
        avg_min_fitness, avg_fitness_after_first, sd_min_fitness,
        avg_generations, relative_early_convergence, area_under_convergence_curve

Outputs (written to --outdir):
- rel_improvements_per_combination_AFI.csv
- rel_improvements_overall_AFI.csv  (includes BOTH Generations Δ% metrics)
- rel_improvements_standard_by_dim_sat_AFI.csv
- rel_improvements_custom_by_sat_AFI.csv
- generations_function_equal_per_function.csv  (per-function averages + Δ% diag)
- GA_comparison_AFI_REC_AUCC.docx (requires python-docx)

Usage:
    python analyze_ga_results_from_csv_v2.py --input _final_results_merged.csv --outdir outputs
"""
import argparse, os, sys, math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# -----------------------------
# Core metric helpers
# -----------------------------
def adjusted_fitness_improvement(f_first, f_base, f_div):
    """AFI (%) = ((f_base - f_div) / (f_first - f_div)) * 100"""
    denom = f_first - f_div
    if denom == 0:
        return np.nan
    return ((f_base - f_div) / denom) * 100

def rel_improvement_pct(base, other):
    """100 * (base - other) / |base| with safe checks (higher = other better if lower is better)."""
    if base == 0 or np.isnan(base) or np.isnan(other):
        return np.nan
    return 100.0 * (base - other) / abs(base)

# -----------------------------
# Plotting Utilities (matplotlib only)
# -----------------------------
def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path

def plot_lines(df, x_col, y_col, hue_col, title, xlabel, ylabel, outpath):
    """Simple line chart with one figure, multiple lines by hue_col (e.g., compare_to)."""
    if df.empty:
        return
    plt.figure()
    for key in sorted(df[hue_col].dropna().unique()):
        sub = df[df[hue_col] == key].sort_values(x_col)
        if sub.empty:
            continue
        plt.plot(sub[x_col].astype(int), sub[y_col], marker='o', label=str(key))
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.legend()
    plt.tight_layout()
    plt.savefig(outpath, dpi=200)
    plt.close()

# -----------------------------
# Main computations
# -----------------------------
def compute_pairwise(df):
    """Compute AFI, REC (corrected), AUCC, generations, ΔSD for Diversity vs baselines per (exp, dim, sat)."""
    rows = []
    group_cols = ["experiment","number_of_variables","saturation_generations"]
    required = {
        "strategy","avg_min_fitness","avg_fitness_after_first","sd_min_fitness",
        "avg_generations","relative_early_convergence","area_under_convergence_curve"
    }
    miss = required - set(df.columns)
    if miss:
        raise KeyError(f"Input file is missing required columns: {sorted(miss)}")

    for (exp, nvars, sat), sub in df.groupby(group_cols):
        means = sub.groupby("strategy").agg(
            fmin=("avg_min_fitness","mean"),
            ffirst=("avg_fitness_after_first","mean"),
            sd=("sd_min_fitness","mean"),
            gens=("avg_generations","mean"),
            rec=("relative_early_convergence","mean"),
            aucc=("area_under_convergence_curve","mean")
        )
        if "diversity mutation" not in means.index:
            continue  # need diversity to compare against
        div = means.loc["diversity mutation"]

        for baseline in ["adaptive mutation","random mutation"]:
            if baseline not in means.index:
                continue
            base = means.loc[baseline]

            afi = adjusted_fitness_improvement(div.ffirst, base.fmin, div.fmin)
            aucc_imp = rel_improvement_pct(base.aucc, div.aucc)
            rec_imp  = rel_improvement_pct(base.rec,  div.rec)   # lower REC better ⇒ positive Δ% means Diversity earlier
            gens_imp = rel_improvement_pct(base.gens, div.gens)  # lower gens better
            dsd      = rel_improvement_pct(base.sd,   div.sd)    # positive ⇒ Diversity more stable

            rows.append({
                "experiment": exp,
                "number_of_variables": nvars,
                "saturation": sat,
                "compare_to": baseline,
                "adjusted_fitness_improvement_%": afi,
                "aucc_improvement_%": aucc_imp,
                "rec_improvement_%": rec_imp,
                "generations_improvement_%": gens_imp,  # classic, slice-weighted
                "stability_delta_SD": dsd
            })
    return pd.DataFrame(rows)

def compute_function_equal_generations(df):
    """
    Treat each function equally regardless of dimensionality/saturation:
    1) Average avg_generations per (experiment, strategy) across all dims/sats.
    2) For each function, compute Δ% that Diversity reduces generations vs baseline:
         Δ%_func = 100 * (G_base_func - G_div_func) / G_base_func
    3) Report per-function table + overall average across functions (unweighted).
    Returns:
        per_func (DataFrame): columns [experiment, adaptive, diversity, random, d_vs_adaptive_%, d_vs_random_%]
        overall  (DataFrame): rows compare_to with mean Δ% across functions.
    """
    required = {"experiment","strategy","avg_generations"}
    if not required.issubset(df.columns):
        raise KeyError(f"Missing columns for function-equal generations: {sorted(required - set(df.columns))}")

    avg_func = (df.groupby(["experiment","strategy"])["avg_generations"]
                  .mean().reset_index())
    pivot = avg_func.pivot(index="experiment", columns="strategy", values="avg_generations").reset_index()

    needed = {"adaptive mutation","diversity mutation","random mutation"}
    have_all = pivot.dropna(subset=list(needed)).rename(columns={
        "adaptive mutation":"adaptive",
        "diversity mutation":"diversity",
        "random mutation":"random"
    })

    def rel(base, other):
        return np.nan if (pd.isna(base) or base==0 or pd.isna(other)) else 100.0*(base - other)/base

    have_all["d_vs_adaptive_%"] = have_all.apply(lambda r: rel(r["adaptive"], r["diversity"]), axis=1)
    have_all["d_vs_random_%"]   = have_all.apply(lambda r: rel(r["random"],   r["diversity"]), axis=1)

    overall = pd.DataFrame([
        {"compare_to":"adaptive mutation", "generations_improvement_function_equal_%": have_all["d_vs_adaptive_%"].mean()},
        {"compare_to":"random mutation",   "generations_improvement_function_equal_%": have_all["d_vs_random_%"].mean()},
    ])
    return have_all, overall

# -----------------------------
# Reporting
# -----------------------------
def save_standard_plots(std_agg, outdir):
    charts_dir = _ensure_dir(os.path.join(outdir, "charts_standard"))
    metrics = [
        ("adjusted_fitness_improvement_%", "AFI (%)"),
        ("rec_improvement_%", "REC Δ% (higher = faster early convergence)"),
        ("aucc_improvement_%", "AUCC Δ%"),
        ("generations_improvement_%", "Generations Δ% (slice-weighted)"),
        ("stability_delta_SD", "SD Δ% (positive = Diversity more stable)")
    ]
    if std_agg.empty:
        return
    for nvars in sorted(std_agg["number_of_variables"].dropna().unique()):
        sub_dim = std_agg[std_agg["number_of_variables"] == nvars]
        for metric, label in metrics:
            title = f"Standard — {int(nvars)} vars: {label} vs Saturation"
            fname = f"standard_{int(nvars)}vars_{metric.replace('%','pct').replace(' ','_')}.png"
            outpath = os.path.join(charts_dir, fname)
            plot_lines(sub_dim, x_col="saturation", y_col=metric, hue_col="compare_to",
                       title=title, xlabel="Saturation (no-improvement generations)", ylabel=label, outpath=outpath)

def save_custom_plots(cust_agg, outdir):
    charts_dir = _ensure_dir(os.path.join(outdir, "charts_custom"))
    metrics = [
        ("adjusted_fitness_improvement_%", "AFI (%)"),
        ("rec_improvement_%", "REC Δ% (higher = faster early convergence)"),
        ("aucc_improvement_%", "AUCC Δ%"),
        ("generations_improvement_%", "Generations Δ% (slice-weighted)"),
        ("stability_delta_SD", "SD Δ% (positive = Diversity more stable)")
    ]
    if cust_agg.empty:
        return
    for metric, label in metrics:
        title = f"Custom (pooled): {label} vs Saturation"
        fname = f"custom_{metric.replace('%','pct').replace(' ', '_')}.png"
        outpath = os.path.join(charts_dir, fname)
        plot_lines(cust_agg, x_col="saturation", y_col=metric, hue_col="compare_to",
                   title=title, xlabel="Saturation (no-improvement generations)", ylabel=label, outpath=outpath)

def write_word_report(overall_df, std_agg, cust_agg, outpath):
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Skipping Word report.", file=sys.stderr)
        return

    doc = Document()
    doc.add_heading("GA Strategy Comparison with AFI (Adjusted Fitness), REC & AUCC", level=1)
    doc.add_paragraph(
        "AFI accounts for the improvement margin from first-generation average fitness down to the diversity minimum. "
        "REC and AUCC are interpreted as lower-is-better and converted to Δ% so that higher values favor the Diversity mutation. "
        "Generations Δ% is shown in two ways: classic slice-weighted and function-equal (each function contributes one vote). "
        "Stability is SD Δ% = 100×(SD_base − SD_div)/SD_base; positive values indicate lower variability with Diversity."
    )

    # Overall
    doc.add_heading("Overall (all functions, all settings)", level=2)
    t = doc.add_table(rows=1, cols=7)
    hdr = t.rows[0].cells
    hdr[0].text = "Compared To"
    hdr[1].text = "AFI (%)"
    hdr[2].text = "AUCC Δ%"
    hdr[3].text = "REC Δ%"
    hdr[4].text = "Generations Δ% (slice-weighted)"
    hdr[5].text = "Generations Δ% (function-equal)"
    hdr[6].text = "SD Δ%"
    for _, r in overall_df.iterrows():
        rw = t.add_row().cells
        rw[0].text = str(r["compare_to"])
        rw[1].text = f"{r['adjusted_fitness_improvement_%']:.2f}"
        rw[2].text = f"{r['aucc_improvement_%']:.2f}"
        rw[3].text = f"{r['rec_improvement_%']:.2f}"
        rw[4].text = f"{r['generations_improvement_%']:.2f}"
        rw[5].text = f"{r.get('generations_improvement_function_equal_%', float('nan')):.2f}"
        rw[6].text = f"{r['stability_delta_SD']:.2f}"

    # Standard Benchmarks (by dimension & saturation)
    doc.add_heading("Standard Benchmarks (by dimension & saturation)", level=2)
    for nvars in [2,3,5,7]:
        sub = std_agg[std_agg["number_of_variables"]==nvars]
        doc.add_heading(f"{nvars} variables", level=3)
        if sub.empty:
            doc.add_paragraph("No data.")
            continue
        t = doc.add_table(rows=1, cols=7)
        hdr = t.rows[0].cells
        hdr[0].text = "Compared To"; hdr[1].text = "Saturation"; hdr[2].text = "AFI (%)"; hdr[3].text = "AUCC Δ%"; hdr[4].text = "REC Δ%"; hdr[5].text = "Generations Δ%"; hdr[6].text = "SD Δ%"
        for _, r in sub.iterrows():
            rw = t.add_row().cells
            rw[0].text = r["compare_to"]
            rw[1].text = str(int(r["saturation"]))
            rw[2].text = f"{r['adjusted_fitness_improvement_%']:.2f}"
            rw[3].text = f"{r['aucc_improvement_%']:.2f}"
            rw[4].text = f"{r['rec_improvement_%']:.2f}"
            rw[5].text = f"{r['generations_improvement_%']:.2f}"
            rw[6].text = f"{r['stability_delta_SD']:.2f}"

    # Custom Trigonometric Functions (pooled by saturation)
    doc.add_heading("Custom Trigonometric Functions (pooled by saturation)", level=2)
    t = doc.add_table(rows=1, cols=7)
    hdr = t.rows[0].cells
    hdr[0].text = "Compared To"; hdr[1].text = "Saturation"; hdr[2].text = "AFI (%)"; hdr[3].text = "AUCC Δ%"; hdr[4].text = "REC Δ%"; hdr[5].text = "Generations Δ%"; hdr[6].text = "SD Δ%"
    for _, r in cust_agg.iterrows():
        rw = t.add_row().cells
        rw[0].text = r["compare_to"]
        rw[1].text = str(int(r["saturation"]))
        rw[2].text = f"{r['adjusted_fitness_improvement_%']:.2f}"
        rw[3].text = f"{r['aucc_improvement_%']:.2f}"
        rw[4].text = f"{r['rec_improvement_%']:.2f}"
        rw[5].text = f"{r['generations_improvement_%']:.2f}"
        rw[6].text = f"{r['stability_delta_SD']:.2f}"

    doc.save(outpath)

# -----------------------------
# Driver
# -----------------------------
def generate_results(input_path: str, outdir: str):
    os.makedirs(outdir, exist_ok=True)
    df = pd.read_csv(input_path)

    # Pairwise comparisons (per exp/dim/sat)
    rel_all = compute_pairwise(df)

    # Save per-combination
    per_combo_path = os.path.join(outdir, "rel_improvements_per_combination_AFI.csv")
    rel_all.to_csv(per_combo_path, index=False)

    # Overall averages (slice-weighted across exp/dim/sat)
    overall = (rel_all.groupby("compare_to")[
                   ["adjusted_fitness_improvement_%", "aucc_improvement_%", "rec_improvement_%",
                    "generations_improvement_%", "stability_delta_SD"]
               ].mean().reset_index())

    # Function-equal Generations Δ% (each function counts once)
    per_func, func_overall = compute_function_equal_generations(df)
    func_per_path = os.path.join(outdir, "generations_function_equal_per_function.csv")
    per_func.to_csv(func_per_path, index=False)

    # Merge function-equal column into overall
    overall = overall.merge(func_overall, on="compare_to", how="left")

    # Save overall table
    overall_path = os.path.join(outdir, "rel_improvements_overall_AFI.csv")
    overall.to_csv(overall_path, index=False)

    # Standard vs Custom aggregations for plots/tables
    standard_funcs = [
        "Sphere Function", "Rosenbrock Function", "Ackley Function", "Rastrigin Function",
        "Griewank Function", "Beale Function", "Himmelblau Function", "Booth Function",
        "Styblinski Tang Function"
    ]
    custom_funcs = [
        "Highly Coupled Trigonometric Function",
        "Moderately Coupled Trigonometric Function",
        "Separated Trig Arithmetic Function"
    ]

    rel_std = rel_all[rel_all["experiment"].isin(standard_funcs) & rel_all["number_of_variables"].isin([2, 3, 5, 7])]
    std_agg = (rel_std.groupby(["number_of_variables", "saturation", "compare_to"])[
                   ["adjusted_fitness_improvement_%", "aucc_improvement_%", "rec_improvement_%",
                    "generations_improvement_%", "stability_delta_SD"]
               ].mean().reset_index().sort_values(["number_of_variables", "compare_to", "saturation"]))
    std_agg_path = os.path.join(outdir, "rel_improvements_standard_by_dim_sat_AFI.csv")
    std_agg.to_csv(std_agg_path, index=False)

    rel_cust = rel_all[rel_all["experiment"].isin(custom_funcs)]
    cust_agg = (rel_cust.groupby(["saturation", "compare_to"])[
                    ["adjusted_fitness_improvement_%", "aucc_improvement_%", "rec_improvement_%",
                     "generations_improvement_%", "stability_delta_SD"]
                ].mean().reset_index().sort_values(["compare_to", "saturation"]))
    cust_agg_path = os.path.join(outdir, "rel_improvements_custom_by_sat_AFI.csv")
    cust_agg.to_csv(cust_agg_path, index=False)

    # Charts
    try:
        save_standard_plots(std_agg, outdir)
        save_custom_plots(cust_agg, outdir)
    except Exception as e:
        print("Warning: could not generate plots:", e, file=sys.stderr)

    # Word report
    word_path = os.path.join(outdir, "GA_comparison_AFI_REC_AUCC.docx")
    try:
        write_word_report(overall, std_agg, cust_agg, word_path)
    except Exception as e:
        print("Warning: could not write Word report:", e, file=sys.stderr)

    print("Saved:")
    print(" -", per_combo_path)
    print(" -", overall_path)
    print(" -", std_agg_path)
    print(" -", cust_agg_path)
    print(" -", func_per_path)
    print(" -", word_path)

def main():
    parser = argparse.ArgumentParser(description="Analyze GA results from CSV with AFI/REC/AUCC.")
    parser.add_argument("--input", default="_final_results_merged.csv", help="Path to GA results CSV.")
    parser.add_argument("--outdir", default=".", help="Output directory for CSVs and DOCX.")
    args = parser.parse_args()
    generate_results(args.input, args.outdir)

if __name__ == "__main__":
    main()
